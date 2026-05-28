from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOCX_PATH = "Relatorio_TP2_Academico.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.167
    p.add_run(text)


def add_section(doc, title, paragraphs):
    doc.add_heading(title, level=1)
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for style_name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Relatorio Tecnico - Trabalho Pratico 2")
run.bold = True
run.font.name = "Calibri"
run.font.size = Pt(18)
run.font.color.rgb = RGBColor.from_string("0B2545")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Autenticacao e Integracao de Dados - Sistema Academico").italic = True

metadata = doc.add_table(rows=4, cols=2)
metadata.autofit = False
rows = [
    ("Disciplina", "Programacao WEB I"),
    ("Curso", "TADS - 3o semestre"),
    ("Tecnologias", "ASP.NET Core .NET 10, Entity Framework Core e MySQL"),
    ("Entrega", "Implementacao WEB funcional com relatorio PDF"),
]
for row, values in zip(metadata.rows, rows):
    for index, value in enumerate(values):
        cell = row.cells[index]
        cell.text = value
        set_cell_width(cell, Inches(1.875 if index == 0 else 4.625))
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if index == 0:
            set_cell_shading(cell, "F2F4F7")
            cell.paragraphs[0].runs[0].bold = True

doc.add_heading("1. Descricao da arquitetura adotada", level=1)
doc.add_paragraph(
    "O sistema foi organizado no padrao MVC do ASP.NET Core, separando controllers, models, views, repositories, servicos e a camada de dados. "
    "Os controllers tratam o fluxo HTTP e a navegacao, as views renderizam os formularios e listagens, os models representam as entidades do dominio e os repositories concentram o acesso ao banco."
)
add_bullet(doc, "Controllers: AlunoController, ProfessorController, DisciplinaController e AuthController.")
add_bullet(doc, "Repositories: camada responsavel pelas operacoes de consulta, criacao, atualizacao e exclusao no banco de dados.")
add_bullet(doc, "Data: AppDbContext centraliza o mapeamento das entidades para o Entity Framework Core.")
add_bullet(doc, "Services: PasswordHasher encapsula a geracao e validacao de hashes de senha.")

doc.add_heading("2. Principais decisoes de implementacao", level=1)
add_bullet(doc, "Foi utilizado cookie authentication para controlar a sessao do usuario autenticado.")
add_bullet(doc, "As paginas internas receberam [Authorize], garantindo acesso apenas apos login.")
add_bullet(doc, "A string de conexao e o usuario padrao foram centralizados no appsettings.json.")
add_bullet(doc, "A criacao e atualizacao do usuario padrao foi movida para UsuarioRepository, mantendo CRUD concentrado nos repositories.")
add_bullet(doc, "As senhas sao armazenadas somente como hash PBKDF2, nunca em texto plano na tabela Usuarios.")

doc.add_heading("3. Funcionamento da autenticacao", level=1)
doc.add_paragraph(
    "A primeira pagina do sistema oferece acesso a area de login. O usuario informa email e senha, e o AuthController consulta o UsuarioRepository pelo email informado. "
    "Quando o usuario existe e esta ativo, o PasswordHasher compara a senha digitada com o hash armazenado no banco. Se a validacao for bem-sucedida, o sistema cria claims de identificacao e inicia a autenticacao por cookie."
)
doc.add_paragraph(
    "O logout remove o cookie de autenticacao e limpa a sessao. Dessa forma, as paginas de alunos, professores e disciplinas voltam a exigir login."
)

doc.add_heading("4. Estrutura do banco de dados", level=1)
doc.add_paragraph("O banco MySQL e acessado por Entity Framework Core. O modelo principal possui as seguintes tabelas:")
table = doc.add_table(rows=1, cols=3)
table.autofit = False
headers = ["Tabela", "Finalidade", "Campos principais"]
for idx, header in enumerate(headers):
    cell = table.rows[0].cells[idx]
    cell.text = header
    set_cell_shading(cell, "F2F4F7")
    set_cell_margins(cell)
    cell.paragraphs[0].runs[0].bold = True
widths = [Inches(1.4), Inches(2.2), Inches(2.9)]
data = [
    ("Usuarios", "Controle de acesso", "Id, Nome, Email, PasswordHash, Perfil, Ativo"),
    ("Pessoas", "Dados comuns de pessoas", "Id, Nome, CPF, DataNascimento"),
    ("Alunos", "Dados especificos de discentes", "Id, Matricula, Curso"),
    ("Professores", "Dados especificos de docentes", "Id, Siape, Area"),
    ("Disciplinas", "Componentes curriculares", "Id, Nome, CargaHoraria, Periodo, ProfessorId"),
    ("AlunoDisciplina", "Relacionamento N:N", "AlunosId, DisciplinasId"),
]
for values in data:
    cells = table.add_row().cells
    for idx, value in enumerate(values):
        cells[idx].text = value
        set_cell_width(cells[idx], widths[idx])
        set_cell_margins(cells[idx])
        cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc.add_heading("5. Dificuldades encontradas", level=1)
add_bullet(doc, "A exclusao direta com ExecuteDeleteAsync nao funcionava nas entidades Aluno e Professor por causa do mapeamento TPT.")
add_bullet(doc, "O executavel em execucao bloqueava a recompilacao no Windows, exigindo parar o processo antes do build.")
add_bullet(doc, "Foi necessario evitar dados fixos no codigo, especialmente credenciais de usuario e seed de autenticacao.")

doc.add_heading("6. Estrategias de resolucao", level=1)
add_bullet(doc, "As exclusoes passaram a carregar a entidade pelo EF, limpar vinculos quando necessario e remover com SaveChangesAsync.")
add_bullet(doc, "Os botoes de exclusao foram trocados para formularios POST com token antifalsificacao.")
add_bullet(doc, "O usuario padrao passou a ser configurado em appsettings.json e garantido pelo UsuarioRepository na inicializacao.")
add_bullet(doc, "A validacao foi feita com dotnet build e busca por ocorrencias de seed fixo, credenciais no codigo e ExecuteDeleteAsync.")

doc.add_heading("7. Observacao sobre uso de IA", level=1)
doc.add_paragraph(
    "A inteligencia artificial foi utilizada como apoio para revisao, organizacao e ajustes tecnicos. O estudante deve revisar o codigo, executar o sistema e estar preparado para explicar a arquitetura, o fluxo de autenticacao, o uso dos repositories e a estrutura do banco."
)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Sistema Academico - TP2 - Programacao WEB I").font.size = Pt(9)

doc.save(DOCX_PATH)
print(DOCX_PATH)
